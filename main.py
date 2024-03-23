import sys
import math
from PySide6.QtWidgets import QApplication, QLabel, QPushButton, QFileDialog, QVBoxLayout, QWidget,QTableWidget,QTableWidgetItem,QMainWindow,QCheckBox
from PySide6.QtCore import Qt
from PySide6 import QtCharts,QtGui,QtWidgets
from PySide6.QtGui import QPainter
from openpyxl import load_workbook
from docx import Document
import numpy as np
from scipy import stats
from scipy.stats import t

def pearson_chi_square(observed_freq, expected_freq,r):
    # Проверяем, что размерности массивов совпадают


    # Вычисляем наблюдаемую и ожидаемую общие частоты
    observed_total = np.sum(observed_freq)
    expected_total = np.sum(expected_freq)

    # Вычисляем статистику Хи-квадрат
    chi_square_statistic = np.sum((observed_freq - expected_freq) ** 2 / expected_freq)

    # Вычисляем степени свободы
    degrees_of_freedom = observed_freq.size - r


    return chi_square_statistic, degrees_of_freedom

def find_critical_point(alpha, degrees_of_freedom):
    # Найти критическую точку
    critical_point = stats.chi2.ppf(1 - alpha, degrees_of_freedom)
    return critical_point
data={(0,10):393,(10,20):238,(20,30):144,(30,40):87,(40,50):53,(50,60):32,(60,70):19}

def test_e(data):
    xinterval=[]
    ni=[]
    for key,val in data.items():
        xinterval.append(key)
        ni.append(val)
    xi = [int((x[0]+x[1])/2) for x in xinterval]
    #ni = [365, 245, 150, 100, 70,45,25]
    # Вычисление суммы наблюдаемых частот
    n_total = np.sum(ni)

    # Шаг 1: Вычисление выборочной средней
    x = np.sum(np.array(xi) * np.array(ni)) / n_total
    # Шаг 2: Оценка параметра показательного распределения
    lmbda = 1 / x

    # Шаг 3: Вычисление вероятностей
    p_interval = [math.exp(-lmbda * i[0]) - math.exp(-lmbda * i[1]) for i in xinterval]

    # Шаг 4: Вычисление ожидаемых частот
    n_prime = np.array(p_interval) * n_total
    observed_freq = np.array(ni)
    expected_freq = n_prime
    chi_square,df = pearson_chi_square(observed_freq,expected_freq,2)
    critical_point = find_critical_point(0.05, df)
    if chi_square < critical_point and chi_square>0:
        return ('Данные соответствуют показательному распределению')
    else:
        return ('Данные не соответствуют показательному распределению')

def test_n(data):
    xinterval = []
    ni = []
    for key, val in data.items():
        xinterval.append(key)
        ni.append(val)
    xi = [((x[0] + x[1]) / 2) for x in xinterval]
    # ni = [365, 245, 150, 100, 70,45,25]
    # Вычисление суммы наблюдаемых частот
    n_total = np.sum(ni)

    # Шаг 1: Вычисление выборочного среднего и выборочной дисперсии
    x_mean = np.sum(np.array(xi) * np.array(ni)) / n_total
    x_var = np.sum(np.array([(xi_i - x_mean) ** 2 for xi_i in xi]) * np.array(ni)) / n_total

    # Шаг 2: Оценка параметров нормального распределения
    mu = x_mean
    sigma = np.sqrt(x_var)
    # Шаг 3: Вычисление вероятностей
    #print((stats.norm.cdf((15-mu)/sigma)-0.5)-(stats.norm.cdf((13-mu)/sigma)-0.5))
    p_interval = [(stats.norm.cdf(xinterval[i][1], loc=mu, scale=sigma)-0.5) - (stats.norm.cdf(xinterval[i][0], loc=mu, scale=sigma)-0.5)  for i in range(len(xi))]

    # Шаг 4: Вычисление ожидаемых частот
    n_prime = np.array(p_interval) * n_total
    n_prime=np.around(n_prime, decimals=0)
    observed_freq = np.array(ni)
    expected_freq = n_prime
    chi_square, df = pearson_chi_square(observed_freq, expected_freq, 3)

    critical_point = find_critical_point(0.05, df)
    if chi_square < critical_point and chi_square>0:
        return ('Данные соответствуют нормальному распределению')
    else:
        return ('Данные не соответствуют нормальному распределению')

def test_u(data):
    xinterval = []
    ni = []
    for key, val in data.items():
        xinterval.append(list(key))
        ni.append(val)
    xi = [((x[0] + x[1]) / 2) for x in xinterval]
    # ni = [365, 245, 150, 100, 70,45,25]
    # Вычисление суммы наблюдаемых частот
    n_total = np.sum(ni)

    # Шаг 1: Вычисление выборочного среднего и выборочной дисперсии
    x_mean = np.sum(np.array(xi) * np.array(ni)) / n_total
    x_var = np.sum(np.array([(xi_i - x_mean) ** 2 for xi_i in xi]) * np.array(ni)) / n_total

    # Шаг 2: Оценка параметров нормального распределения
    mu = x_mean
    sigma = np.sqrt(x_var)

    a=mu-math.sqrt(3)*sigma
    b = mu + math.sqrt(3) * sigma
    xinterval[0][0]=a
    xinterval[-1][1] = b

    # Шаг 3: Вычисление вероятностей
    #print((stats.norm.cdf((15-mu)/sigma)-0.5)-(stats.norm.cdf((13-mu)/sigma)-0.5))
    p_interval = [((x[1]-x[0])/(b-a)) for x in xinterval]

    # Шаг 4: Вычисление ожидаемых частот
    n_prime = np.array(p_interval) * n_total
    n_prime=np.around(n_prime, decimals=0)

    observed_freq = np.array(ni)
    expected_freq = n_prime
    chi_square, df = pearson_chi_square(observed_freq, expected_freq, 3)

    critical_point = find_critical_point(0.05, df)
    if chi_square < critical_point and chi_square>0:
        return('Данные соответствуют равномерному распределению')
    else:
        return('Данные не соответствуют равномерному распределению')

def tooltip_handler(point, state, chart):
    tooltip = QtWidgets.QToolTip
    if state:
        tooltip.showText(chart_view.mapToGlobal(point.toPoint()), str(round(point.y())))
    else:
        tooltip.hideText()

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

def count_numbers_in_intervals(intervals, numbers):
    result = {}

    for interval in intervals:
        start = interval[0]
        end = interval[1]
        count = 0

        for number in numbers:
            if start < number <= end:
                count += 1

        result[tuple(interval)] = count

    return result

class DragDropLabel(QLabel):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setText('Перетащите файл XLSX сюда')

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        files = [url.toLocalFile() for url in event.mimeData().urls()]
        if len(files) > 0:
            file_path = files[0]
            if file_path.endswith('.xlsx') or file_path.endswith('.docx'):
                self.setText(f'Выбран файл: {file_path}')
                self.file_path = file_path
            else:
                self.setText('Неверное расширение файла. Пожалуйста, выберите файл XLSX.')

def open_file_dialog():
    file_dialog = QFileDialog()
    file_dialog.setNameFilters(["Word Files (*.docx)", "Excel Files (*.xlsx)"])
    file_dialog.setFileMode(QFileDialog.ExistingFile)
    if file_dialog.exec():
        file_path = file_dialog.selectedFiles()[0]
        if file_path.endswith('.xlsx') or file_path.endswith('.docx'):
            label.setText(f'Выбран файл: {file_path}')
            label.file_path = file_path
        else:
            label.setText('Неверное расширение файла. Пожалуйста, выберите файл XLSX.')
table = None
table2=None
text1=None
chart_view=None
chart_view2=None
checkbox = None

def data_check(data):
        global table, table2, text1, chart_view, chart_view2,checkbox
        maxX = max(data)
        minX = min(data)
        print(sum(data)/len(data))
        dx = maxX - minX
        r = 1 + 3.2 * math.log10(len(data))
        r = int(r + (0.5 if r > 0 else -0.5))
        rlenth = dx / r
        start = minX - rlenth / 2
        i = 0
        intervals = [start]
        while start < maxX:
            intervals.append(start + rlenth)
            start += rlenth
        pairs = []
        for i, v in enumerate(intervals):
            if i != 0:
                pairs.append([intervals[i - 1], v])
        tablearr = count_numbers_in_intervals(pairs, data)
        table2arr = {}
        for key, val in tablearr.items():
            table2arr[(key[0] + key[1]) / 2] = val

        # Создание таблицы
        table = QTableWidget()
        table.setRowCount(2)
        table.setColumnCount(len(tablearr) + 1)
        table.verticalHeader().setVisible(False)
        table.horizontalHeader().setVisible(False)
        table.setItem(0, 0, QTableWidgetItem('Xi'))
        table.setItem(1, 0, QTableWidgetItem('Ni'))
        i = 0
        # Заполнение таблицы данными из массивов
        for key, val in tablearr.items():
            item1 = QTableWidgetItem(str(round(key[0], 2)) + ' - ' + str(round(key[1], 2)))
            item1.setFlags(item1.flags() ^ Qt.ItemIsEditable)
            item2 = QTableWidgetItem(str(val))
            item2.setFlags(item2.flags() ^ Qt.ItemIsEditable)
            table.setItem(0, i + 1, item1)
            table.setItem(1, i + 1, item2)
            i += 1

        # Установка размера ячеек
        table.resizeColumnsToContents()
        table.resizeRowsToContents()

        # Создание таблицы
        table2 = QTableWidget()
        table2.setRowCount(3)
        table2.setColumnCount(len(tablearr) + 1)
        table2.verticalHeader().setVisible(False)
        table2.horizontalHeader().setVisible(False)
        table2.setItem(0, 0, QTableWidgetItem('Xi'))
        table2.setItem(1, 0, QTableWidgetItem('Ni'))
        table2.setItem(2, 0, QTableWidgetItem('Wi'))
        i = 0

        # Заполнение таблицы данными из массивов
        for key, val in table2arr.items():
            item1 = QTableWidgetItem(str(round(key, 2)))
            item1.setFlags(item1.flags() ^ Qt.ItemIsEditable)
            item2 = QTableWidgetItem(str(val))
            item2.setFlags(item2.flags() ^ Qt.ItemIsEditable)
            item3 = QTableWidgetItem(str(val / len(data)))
            # table2arr[key]=val/len(data)
            item2.setFlags(item2.flags() ^ Qt.ItemIsEditable)
            table2.setItem(0, i + 1, item1)
            table2.setItem(1, i + 1, item2)

            table2.setItem(2, i + 1, item3)
            i += 1
        x2 = 0
        xv = sum(data)/len(data)
        for key, val in table2arr.items():
            x2 += (key ** 2) * (val / len(data))
        print(x2,xv)
        Dv = x2 - xv ** 2
        sigma = Dv ** (1 / 2)
        V = (sigma / xv) * 100
        S = len(data) / (len(data) - 1) * Dv
        if len(data) > 50:
            text1 = QLabel('Средняя выборочная:' + str(round(xv,2)) + '\n' + 'Дисперсия:' + str(round(Dv,2)) + '\n' + 'Стандарт:' + str(round(sigma,2)) + '\n' + 'Размах варьирования равен ' + str(round(dx,2)) + '\n' + 'Коэф. вариации равен ' + str(round(V,0)) + '\n' + 'Исправленная выборочная дисперсия равна' + str(round(S,2)) + '\n' + test_n(tablearr) + '\n' + test_u(tablearr) + '\n' + test_e(tablearr))
        else:
            text1 = QLabel('Средняя выборочная:' + str(round(xv,2)) + '\n' + 'Дисперсия:' + str(round(Dv,2)) + '\n' + 'Стандарт:' + str(round(sigma,2)) + '\n' + 'Размах варьирования равен ' + str(round(dx,2)) + '\n' + 'Коэф. вариации равен ' + str(round(V,0)) + '\n' + 'Исправленная выборочная дисперсия равна' + str(round(S,2)) + '\n' + 'Для проверки гипотезы о виде распределения изучаемого признака, объем выборки является недостаточным.(требование репрезентативности выборки не выполняется)')
        # Добавление таблицы на форму
        table.setFixedHeight(90)
        table2.setFixedHeight(90)
        horizontal_layout.addWidget(table)
        horizontal_layout.addWidget(table2)
        font = QtGui.QFont("Arial", 16)  # Создаем экземпляр объекта QFont с размером шрифта 16
        text1.setFont(font)
        layout.addWidget(text1)
        # layout.addWidget(text3)
        # layout.addWidget(text4)
        # layout.addWidget(text5)
        # Установка размера ячеекsigma
        table2.resizeColumnsToContents()
        table2.resizeRowsToContents()
        chart_view = QtCharts.QChartView()
        chart_view.setRenderHint(QPainter.Antialiasing)
        chart_view.setFixedHeight(300)
        horizontal_layout2.addWidget(chart_view)
        chart_view2 = QtCharts.QChartView()
        chart_view2.setRenderHint(QPainter.Antialiasing)
        chart_view2.setFixedHeight(300)
        horizontal_layout2.addWidget(chart_view2)
        x_values = list(table2arr.keys())
        y_values = list(table2arr.values())
        series = QtCharts.QLineSeries()
        series.hovered.connect(lambda point, state: tooltip_handler(point, state, chart))
        for x, y in zip(x_values, y_values):
            series.append(x, y)

        x_values = list(tablearr.keys())
        y_values = list(tablearr.values())
        series2 = QtCharts.QBarSeries()
        for x, y in zip(x_values, y_values):
            barset = QtCharts.QBarSet(str(round(x[0], 2)) + '-' + str(round(x[1], 2)))
            barset.append(y / len(data))
            series2.append(barset)
        chart2 = QtCharts.QChart()
        chart2.addSeries(series2)
        chart2.createDefaultAxes()
        chart_view2.setChart(chart2)
        chart_view2.setRenderHint(QtGui.QPainter.Antialiasing)
        # Создаем график и добавляем серию
        chart = QtCharts.QChart()

        chart.addSeries(series)
        chart.createDefaultAxes()
        chart_view.setChart(chart)
        chart_view.setRenderHint(QtGui.QPainter.Antialiasing)
        # Отображение формы
        window.show()


        label.setText('Значения успешно записаны в массивы.')


def check_numbers():
    global table,table2,text1,chart_view,chart_view2,checkbox
    try:
        if text1:
            layout.removeWidget(text1)
            text1.deleteLater()
        if chart_view:
            layout.removeWidget(chart_view)
            chart_view.deleteLater()
        if chart_view2:
            layout.removeWidget(chart_view2)
            chart_view2.deleteLater()
        if table:
            layout.removeWidget(table)
            table.deleteLater()
        if table2:
            layout.removeWidget(table2)
            table2.deleteLater()
        if checkbox:
            print("---------------------------------------------------------------")
            layout.removeWidget(checkbox)
            checkbox.deleteLater()
    except RuntimeError:
        pass
    file_path = label.file_path
    data = []
    data2 = []
    if file_path.endswith('.xlsx'):
        workbook = load_workbook(file_path,data_only=True)
        sheet = workbook.active
        for row in sheet.iter_rows():
            if len(row) > 0:
                value1 = row[0].value
                data.append(value1)

            if len(row) > 1:
                value2 = row[1].value
                data2.append(value2)
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        text = ""
        text2=""
        if len(doc.paragraphs)==2:
            text += doc.paragraphs[0].text
            text2 += doc.paragraphs[1].text
        elif  len(doc.paragraphs)==1:
            text += doc.paragraphs[0].text
        print(text)
        try:
            data = [int(num) for num in text.split()]
            data2 = [int(num) for num in text2.split()]

        except ValueError:
            label.setText("Ошибка! В строке должны быть только числа, разделенные пробелами.")
            return None
    print(data)
    print(data2)
    if len(data) != 0:
        data_check(data)
    else:
        label.setText('Неверный формат данных')


def check_numbers2():
    global table, table2, text1, chart_view, chart_view2,checkbox
    try:
        if text1:
            layout.removeWidget(text1)
            text1.deleteLater()
        if chart_view:
            layout.removeWidget(chart_view)
            chart_view.deleteLater()
        if chart_view2:
            layout.removeWidget(chart_view2)
            chart_view2.deleteLater()
        if table:
            layout.removeWidget(table)
            table.deleteLater()
        if table2:
            layout.removeWidget(table2)
            table2.deleteLater()
        if checkbox:
            print("---------------------------------------------------------------")
            layout.removeWidget(checkbox)
            checkbox.deleteLater()
    except RuntimeError:
        pass
    file_path = label.file_path
    data = []
    data2 = []
    if file_path.endswith('.xlsx'):
        workbook = load_workbook(file_path, data_only=True)
        sheet = workbook.active
        for row in sheet.iter_rows():
            if len(row) > 0:
                value1 = row[0].value
                data.append(value1)

            if len(row) > 1:
                value2 = row[1].value
                data2.append(value2)
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        text = ""
        text2 = ""
        if len(doc.paragraphs) == 2:
            text += doc.paragraphs[0].text
            text2 += doc.paragraphs[1].text
        elif len(doc.paragraphs) == 1:
            text += doc.paragraphs[0].text
        print(text)
        try:
            data = [int(num) for num in text.split()]
            data2 = [int(num) for num in text2.split()]

        except ValueError:
            label.setText("Ошибка! В строке должны быть только числа, разделенные пробелами.")
            return None
    print(data)
    print(data2)
    if len(data2) != 0:
        data_check(data2)
    else:
        label.setText('Неверный формат данных')

def correlation():
    global text1,chart_view,checkbox
    try:
        if text1:
            layout.removeWidget(text1)
            text1.deleteLater()
        if chart_view:
            layout.removeWidget(chart_view)
            chart_view.deleteLater()
        if chart_view2:
            layout.removeWidget(chart_view2)
            chart_view2.deleteLater()
        if table:
            layout.removeWidget(table)
            table.deleteLater()
        if table2:
            layout.removeWidget(table2)
            table2.deleteLater()
        if checkbox:
            print("---------------------------------------------------------------")
            layout.removeWidget(checkbox)
            checkbox.deleteLater()
    except RuntimeError:
        pass
    # try:
    #     if table:
    #         layout.removeWidget(chart_view)
    #         chart_view.deleteLater()
    #         layout.removeWidget(chart_view2)
    #         chart_view2.deleteLater()
    #         layout.removeWidget(text1)
    #         text1.deleteLater()
    #         layout.removeWidget(table)
    #         table.deleteLater()
    #         layout.removeWidget(table2)
    #         table2.deleteLater()
    # except RuntimeError:
    #     pass
    file_path = label.file_path
    data = []
    data2 = []
    if file_path.endswith('.xlsx'):
        workbook = load_workbook(file_path, data_only=True)
        sheet = workbook.active
        for row in sheet.iter_rows():
            if len(row) > 0:
                value1 = row[0].value
                data.append(float(value1))

            if len(row) > 1:
                value2 = row[1].value
                data2.append(float(value2))
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        text = ""
        text2 = ""
        if len(doc.paragraphs) == 2:
            text += doc.paragraphs[0].text
            text2 += doc.paragraphs[1].text
        elif len(doc.paragraphs) == 1:
            text += doc.paragraphs[0].text
        print(text)
        try:
            data = [int(num) for num in text.split()]
            data2 = [int(num) for num in text2.split()]

        except ValueError:
            label.setText("Ошибка! В строке должны быть только числа, разделенные пробелами.")
            return None
    if len(data)==len(data2) and len(data)>0:
        # Расчет средних значений
        mean_x = np.mean(data)
        mean_y = np.mean(data2)
        print(mean_y,mean_x)
        # Расчет разностей
        diff_x = data - mean_x
        diff_y = data2 - mean_y
        chart_view = QtCharts.QChartView()
        # Расчет суммы произведений разностей
        sum_diff_product = np.sum(diff_x * diff_y)

        # Расчет суммы квадратов разностей
        sum_diff_x_squared = np.sum(diff_x ** 2)
        sum_diff_y_squared = np.sum(diff_y ** 2)

        # Расчет выборочного коэффициента корреляции
        corr_coefficient = sum_diff_product / np.sqrt(sum_diff_x_squared * sum_diff_y_squared)
        chart = QtCharts.QChart()
        chart.setTitle("Поле опытных точек")
        print(corr_coefficient)
        series = QtCharts.QLineSeries()
        series.setName("Y на X")
        series2 = QtCharts.QLineSeries()
        series2.setName("X на Y")
        # Заполнение данных для графика функции
        for i in range(len(data)):
            calculated_y = corr_coefficient * (np.std(data2) / np.std(data)) * (data[i] - mean_x) + mean_y
            print(data[i],calculated_y)
            series.append(data[i], calculated_y)
        for i in range(len(data2)):
            calculated_x = corr_coefficient * (np.std(data) / np.std(data2)) * (data2[i] - mean_y) + mean_x
            print(calculated_x, data2[i])
            series2.append(calculated_x, data2[i])
        # Добавление серии данных на график
        chart.addSeries(series)
        chart.addSeries(series2)

        # Создание точечных данных и их добавление на график
        scatter_series = QtCharts.QScatterSeries()
        scatter_series.setName("Точки")
        catter = QtCharts.QScatterSeries()
        catter.setName("Центр рассеивания")
        catter.append(mean_x,mean_y)
        chart.addSeries(catter)
        for i in range(len(data)):
            scatter_series.append(data[i], data2[i])

        scatter_series.setMarkerShape(QtCharts.QScatterSeries.MarkerShapeCircle)
        scatter_series.setMarkerSize(10)

        chart.addSeries(scatter_series)

        # Создание осей
        axisX = QtCharts.QValueAxis()
        axisX.setTitleText("X")
        axisX.setRange(0, max(data))
        chart.addAxis(axisX, Qt.AlignBottom)

        axisY = QtCharts.QValueAxis()
        axisY.setTitleText("Y")
        axisY.setRange(0, max(data2))
        chart.addAxis(axisY, Qt.AlignLeft)
        axisX.setGridLineVisible(True)
        axisX.setMinorGridLineVisible(True)
        axisY.setGridLineVisible(True)
        axisY.setMinorGridLineVisible(True)
        axisX.setTickCount(max(data)+1)  # Устанавливает количество основных делений на оси X
        axisX.setMinorTickCount(0)  # Устанавливает количество дополнительных делений на оси X
        axisY.setTickCount(max(data2)+1)  # Устанавливает количество основных делений на оси X
        axisY.setMinorTickCount(0)  # Устанавливает количество дополнительных делений на оси X
        # Назначение осей для серий данных
        series.attachAxis(axisX)
        series.attachAxis(axisY)
        series2.attachAxis(axisX)
        series2.attachAxis(axisY)
        scatter_series.attachAxis(axisX)
        scatter_series.attachAxis(axisY)
        catter.attachAxis(axisX)
        catter.attachAxis(axisY)
        checkbox = QCheckBox('Отображение линий регрессии(тренда)')
        checkbox.stateChanged.connect(lambda state: series.setVisible(state == 0))
        checkbox.stateChanged.connect(lambda state: series2.setVisible(state == 0))

        # добавляем checkbox на форму
        layout.addWidget(checkbox)
        # Отображение графика
        chart_view.setChart(chart)
        chart_view.setMaximumHeight(900)
        layout.addWidget(chart_view)
        alfa = 1. - 0.95
        print(stats.t.ppf(1. - alfa / 2., len(data)-2))
        расстояние_к_минус_одному = abs(corr_coefficient - (-1))
        расстояние_к_минус_полуторе = abs(corr_coefficient - (-0.5))
        расстояние_к_нулю = abs(corr_coefficient - 0)
        расстояние_к_полуторе = abs(corr_coefficient - 0.5)
        расстояние_к_одному = abs(corr_coefficient - 1)

        минимальное_расстояние = min(расстояние_к_минус_одному, расстояние_к_минус_полуторе, расстояние_к_нулю,
                                     расстояние_к_полуторе, расстояние_к_одному)

        if минимальное_расстояние == расстояние_к_минус_одному:
            text1=QLabel("Сильная отрицательная линейная корреляционная зависимость"+'\n'+'Rv='+str(corr_coefficient))
            font = QtGui.QFont("Arial", 16)  # Создаем экземпляр объекта QFont с размером шрифта 16
            text1.setFont(font)
            container = QVBoxLayout()
            container.addWidget(text1)
            layout.addLayout(container)
        elif минимальное_расстояние == расстояние_к_минус_полуторе:
            text1=QLabel("Средней силы отрицательная линейная корреляционная зависимость"+'\n'+'Rv='+str(corr_coefficient))
            font = QtGui.QFont("Arial", 16)  # Создаем экземпляр объекта QFont с размером шрифта 16
            text1.setFont(font)
            container = QVBoxLayout()
            container.addWidget(text1)
            layout.addLayout(container)
        elif минимальное_расстояние == расстояние_к_нулю:
            text1=QLabel("Слабая линейная корреляционная зависимость"+'\n'+'Rv='+str(corr_coefficient))
            font = QtGui.QFont("Arial", 16)  # Создаем экземпляр объекта QFont с размером шрифта 16
            text1.setFont(font)
            container = QVBoxLayout()
            container.addWidget(text1)
            layout.addLayout(container)
        elif минимальное_расстояние == расстояние_к_полуторе:
            text1=QLabel("Средней силы положительная линейная корреляционная зависимость"+'\n'+'Rv='+str(corr_coefficient))
            font = QtGui.QFont("Arial", 16)  # Создаем экземпляр объекта QFont с размером шрифта 16
            text1.setFont(font)
            container = QVBoxLayout()
            container.addWidget(text1)
            layout.addLayout(container)
        else:
            text1 = QLabel("Сильная положительная линейная корреляционная зависимость"+'\n'+'Rv='+str(corr_coefficient))
            font = QtGui.QFont("Arial", 16)  # Создаем экземпляр объекта QFont с размером шрифта 16
            text1.setFont(font)
            container = QVBoxLayout()
            container.addWidget(text1)
            layout.addLayout(container)

    else:
        label.setText('Неверный формат входных данных')

def read_excel_values():
    global table, table2, text1, chart_view, chart_view2,checkbox
    try:
        if text1:
            layout.removeWidget(text1)
            text1.deleteLater()
        if chart_view:
            layout.removeWidget(chart_view)
            chart_view.deleteLater()
        if chart_view2:
            layout.removeWidget(chart_view2)
            chart_view2.deleteLater()
        if table:
            layout.removeWidget(table)
            table.deleteLater()
        if table2:
            layout.removeWidget(table2)
            table2.deleteLater()
        if checkbox:
            print("---------------------------------------------------------------")
            layout.removeWidget(checkbox)
            checkbox.deleteLater()
    except RuntimeError:
        pass
    file_path = label.file_path
    wb = load_workbook(file_path)
    sheet = wb.active

    first_row = sheet[1]
    second_row = sheet[2]

    tablearr = {}
    ln=0
    try:
        for first_cell, second_cell in zip(first_row, second_row):
            key = [int(num) for num in str(first_cell.value).split('-')]
            value = second_cell.value
            ln+=second_cell.value
            tablearr[tuple(key)] = value
    except ValueError:
        label.setText('Неверные значения')
    table2arr = {}
    data=[]
    try:
        for key, val in tablearr.items():
            data.append(key[0])
            data.append(key[1])
            table2arr[(key[0] + key[1]) / 2] = val
    except IndexError:
        label.setText('Неверный формат данных')
        return
    dx=max(data)-min(data)

    if table:
        layout.removeWidget(chart_view)
        chart_view.deleteLater()
        layout.removeWidget(chart_view2)
        chart_view2.deleteLater()
        layout.removeWidget(text1)
        text1.deleteLater()
        layout.removeWidget(table)
        table.deleteLater()
        layout.removeWidget(table2)
        table2.deleteLater()

    # Создание таблицы
    table = QTableWidget()
    table.setRowCount(2)
    table.setColumnCount(len(tablearr) + 1)
    table.verticalHeader().setVisible(False)
    table.horizontalHeader().setVisible(False)
    table.setItem(0, 0, QTableWidgetItem('Xi'))
    table.setItem(1, 0, QTableWidgetItem('Ni'))
    i = 0
    # Заполнение таблицы данными из массивов
    for key, val in tablearr.items():
        item1 = QTableWidgetItem(str(round(key[0], 2)) + ' - ' + str(round(key[1], 2)))
        item1.setFlags(item1.flags() ^ Qt.ItemIsEditable)
        item2 = QTableWidgetItem(str(val))
        item2.setFlags(item2.flags() ^ Qt.ItemIsEditable)
        table.setItem(0, i + 1, item1)
        table.setItem(1, i + 1, item2)
        i += 1

    # Установка размера ячеек
    table.resizeColumnsToContents()
    table.resizeRowsToContents()

    # Создание таблицы
    table2 = QTableWidget()
    table2.setRowCount(3)
    table2.setColumnCount(len(tablearr) + 1)
    table2.verticalHeader().setVisible(False)
    table2.horizontalHeader().setVisible(False)
    table2.setItem(0, 0, QTableWidgetItem('Xi'))
    table2.setItem(1, 0, QTableWidgetItem('Ni'))
    table2.setItem(2, 0, QTableWidgetItem('Wi'))
    i = 0

    # Заполнение таблицы данными из массивов
    for key, val in table2arr.items():
        item1 = QTableWidgetItem(str(round(key, 2)))
        item1.setFlags(item1.flags() ^ Qt.ItemIsEditable)
        item2 = QTableWidgetItem(str(val))
        item2.setFlags(item2.flags() ^ Qt.ItemIsEditable)
        item3 = QTableWidgetItem(str(val / ln))
        # table2arr[key]=val/len(data)
        item2.setFlags(item2.flags() ^ Qt.ItemIsEditable)
        table2.setItem(0, i + 1, item1)
        table2.setItem(1, i + 1, item2)
        table2.setItem(2, i + 1, item3)
        i += 1
    x2 = 0
    xv = 0
    for key, val in table2arr.items():
        x2 += (key ** 2) * val / ln
        xv += key * val / ln
    Dv = x2 - xv ** 2
    sigma = Dv ** (1 / 2)
    V = (sigma / xv) * 100
    S = ln / (ln - 1) * Dv
    if ln > 50:
        text1 = QLabel('Средняя выборочная:' + str(round(xv,2)) + '\n' + 'Дисперсия:' + str(round(Dv,2)) + '\n' + 'Ствндарт:' + str(round(sigma,2)) + '\n' + 'Размах варьирования равен ' + str(round(dx,2)) + '\n' + 'Коэф. вариации равен ' + str(round(V,0)) + '\n' + 'Исправленная выборочная дисперсия равна' + str(round(S,2)) + '\n' + test_n(tablearr) + '\n' + test_u(tablearr) + '\n' + test_e(tablearr))
    else:
        text1 = QLabel('Средняя выборочная:' + str(round(xv,2)) + '\n' + 'Дисперсия:' + str(round(Dv,2)) + '\n' + 'Стандарт:' + str( sigma) + '\n' + 'Размах варьирования равен ' + str(round(dx,2)) + '\n' + 'Коэф. вариации равен ' + str(round(V,0)) + '\n' + 'Исправленная выборочная дисперсия равна' + str(round(S,2)) + '\n' + 'Для проверки гипотезы о виде распределения изучаемого признака, объем выборки является недостаточным.(требование репрезентативности выборки не выполняется)')
    # Добавление таблицы на форму
    table.setFixedHeight(90)
    table2.setFixedHeight(90)
    horizontal_layout.addWidget(table)
    horizontal_layout.addWidget(table2)
    font = QtGui.QFont("Arial", 16)  # Создаем экземпляр объекта QFont с размером шрифта 16
    text1.setFont(font)
    layout.addWidget(text1)
    # layout.addWidget(text3)
    # layout.addWidget(text4)
    # layout.addWidget(text5)
    # Установка размера ячеекsigma
    table2.resizeColumnsToContents()
    table2.resizeRowsToContents()
    chart_view = QtCharts.QChartView()
    chart_view.setRenderHint(QPainter.Antialiasing)
    chart_view.setFixedHeight(300)
    horizontal_layout2.addWidget(chart_view)
    chart_view2 = QtCharts.QChartView()
    chart_view2.setRenderHint(QPainter.Antialiasing)
    chart_view2.setFixedHeight(300)
    horizontal_layout2.addWidget(chart_view2)
    x_values = list(table2arr.keys())
    y_values = list(table2arr.values())
    series = QtCharts.QLineSeries()
    series.hovered.connect(lambda point, state: tooltip_handler(point, state, chart))
    for x, y in zip(x_values, y_values):
        series.append(x, y)

    x_values = list(tablearr.keys())
    y_values = list(tablearr.values())
    series2 = QtCharts.QBarSeries()
    for x, y in zip(x_values, y_values):
        barset = QtCharts.QBarSet(str(round(x[0], 2)) + '-' + str(round(x[1], 2)))
        barset.append(y / ln)
        series2.append(barset)
    chart2 = QtCharts.QChart()
    chart2.addSeries(series2)
    chart2.createDefaultAxes()
    chart_view2.setChart(chart2)
    chart_view2.setRenderHint(QtGui.QPainter.Antialiasing)
    # Создаем график и добавляем серию
    chart = QtCharts.QChart()

    chart.addSeries(series)
    chart.createDefaultAxes()
    chart_view.setChart(chart)
    chart_view.setRenderHint(QtGui.QPainter.Antialiasing)
    # Отображение формы
    window.show()
    label.setText('Значения успешно записаны в массивы.')


app = QApplication(sys.argv)

label = DragDropLabel()

label.setMinimumHeight(200)
buttonlayout=QtWidgets.QHBoxLayout()
button2 = QPushButton('Статистический ряд распределения признака(закон распределения)')
button2.clicked.connect(read_excel_values)
button = QPushButton('Значения признака Х')
button.clicked.connect(check_numbers)
button3 = QPushButton('Значения признака У')
button3.clicked.connect(check_numbers2)
button4 = QPushButton('Поле опытных точек')
button4.clicked.connect(correlation)
buttonlayout.addWidget(button)
buttonlayout.addWidget(button3)
buttonlayout.addWidget(button2)
buttonlayout.addWidget(button4)
layout = QVBoxLayout()
layout.addWidget(label)
layout.addLayout(buttonlayout)

horizontal_layout = QtWidgets.QHBoxLayout()
layout.addLayout(horizontal_layout)
horizontal_layout2 = QtWidgets.QHBoxLayout()
layout.addLayout(horizontal_layout2)
container = QtWidgets.QWidget()
container.setLayout(layout)


window = QtWidgets.QMainWindow()
window.showMaximized()
window.setCentralWidget(container)
#window.setFixedSize(300, 200)
window.setWindowTitle('Проверка чисел в файле XLSX')
window.show()

sys.exit(app.exec())